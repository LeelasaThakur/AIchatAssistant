import os
import logging

logger = logging.getLogger(__name__)


def allowed_file(filename: str, allowed_extensions: set) -> bool:
    """Return True if the filename has an extension in the whitelist."""
    return "." in filename and filename.rsplit(".", 1)[1].lower() in allowed_extensions


def validate_mime_type(file_path: str) -> bool:
    """
    Basic MIME-type validation using file signatures (magic bytes).
    Returns True if the file appears to be a safe document/image type.
    """
    SAFE_SIGNATURES = {
        b"%PDF":                 "pdf",
        b"PK\x03\x04":          "docx/zip",
        b"\x89PNG":             "png",
        b"\xff\xd8\xff":        "jpeg",
        b"GIF87a":              "gif",
        b"GIF89a":              "gif",
    }

    try:
        with open(file_path, "rb") as f:
            header = f.read(8)

        # Text files have no magic bytes — check if content is text-like
        if not header:
            return False

        for signature in SAFE_SIGNATURES:
            if header.startswith(signature):
                return True

        # If it's not a known binary format, check if it looks like text
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                f.read(512)
            return True  # readable as UTF-8 text
        except (UnicodeDecodeError, ValueError):
            try:
                with open(file_path, "r", encoding="latin-1") as f:
                    f.read(512)
                return True
            except Exception:
                pass

    except OSError as exc:
        logger.warning("MIME validation failed for %s: %s", file_path, exc)

    return False


def extract_text_from_file(file_path: str) -> str | None:
    """
    Extract raw text from TXT, PDF, and DOCX files.
    Returns None for images or unsupported formats.
    """
    ext = os.path.splitext(file_path)[1].lower()

    try:
        if ext == ".txt":
            try:
                with open(file_path, "r", encoding="utf-8") as fh:
                    return fh.read()
            except UnicodeDecodeError:
                with open(file_path, "r", encoding="latin-1") as fh:
                    return fh.read()

        elif ext == ".pdf":
            from pypdf import PdfReader
            reader = PdfReader(file_path)
            pages = [page.extract_text() for page in reader.pages if page.extract_text()]
            return "\n".join(pages)

        elif ext in (".doc", ".docx"):
            import docx
            document = docx.Document(file_path)
            parts = [p.text for p in document.paragraphs]
            for table in document.tables:
                for row in table.rows:
                    parts.append(" | ".join(cell.text for cell in row.cells))
            return "\n".join(parts)

    except Exception as exc:
        logger.warning("Error parsing file %s: %s", file_path, exc)
        return f"[Error extracting text: {type(exc).__name__}]"

    return None  # image or unsupported type